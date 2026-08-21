import os
import json
import base64
import re
import shutil
import hashlib
import html
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
import streamlit as st
import pandas as pd
from pptx import Presentation
import core
import io
from dotenv import load_dotenv

load_dotenv() # .env 파일이 존재하면 로컬 환경변수로 불러옴

# 기본 키 (금고 st.secrets 또는 .env 환경변수에서 우선 가져오기)
API_KEY_DEFAULT = ""
try:
    API_KEY_DEFAULT = st.secrets.get("OPENAI_API_KEY", "")
except Exception:
    pass

if not API_KEY_DEFAULT:
    API_KEY_DEFAULT = os.environ.get("OPENAI_API_KEY", "")

st.set_page_config(page_title="AI 품질관리 시스템(원고, 스토리보드 검토)", page_icon="✨", layout="wide")

# 상단 여백 최소화
st.markdown(
    """
    <style>
    .block-container,
    .stMainBlockContainer,
    [data-testid="stAppViewBlockContainer"] {
        padding-top: 2rem !important;
        padding-bottom: 2rem !important;
    }
    [data-testid="stHeader"] {
        height: 2.5rem !important;
    }
    [data-testid="stSidebar"] {
        width: 31.5rem !important;
        min-width: 31.5rem !important;
    }
    [data-testid="stSidebar"] > div:first-child {
        width: 31.5rem !important;
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("✨AI 품질관리 시스템(원고, 스토리보드 검토)")

APP_DIR = os.path.dirname(os.path.abspath(__file__))
REVIEW_HISTORY_DIR = os.path.join(APP_DIR, "review_history")
SUPABASE_HISTORY_TABLE = "review_history"
SUPABASE_HISTORY_BUCKET = "review-history"
SUPABASE_SETTINGS_TABLE = "app_settings"
_SUPABASE_CLIENT = None
_SUPABASE_ERROR = ""
KST = ZoneInfo("Asia/Seoul")


def _safe_filename(name):
    base = os.path.basename(name or "문서")
    return re.sub(r'[\\/:*?"<>|]+', "_", base).strip() or "문서"


def _safe_storage_segment(value, fallback="file"):
    text = str(value or fallback)
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("._-")
    digest = hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()[:10]
    if cleaned:
        cleaned = cleaned[:80]
        return f"{digest}_{cleaned}"
    return digest or fallback


def _json_default(value):
    if hasattr(value, "item"):
        return value.item()
    return str(value)


def _read_json(path, default=None):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


def _write_json(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, default=_json_default)


def _now_kst():
    return datetime.now(KST)


def _now_kst_iso():
    return _now_kst().isoformat(timespec="seconds")


def _parse_datetime_value(value):
    if not value:
        return None
    try:
        text = str(value)
        if text.endswith("Z"):
            text = text[:-1] + "+00:00"
        dt = datetime.fromisoformat(text)
        if dt.tzinfo is None:
            return dt.replace(tzinfo=KST)
        return dt.astimezone(KST)
    except Exception:
        return None


def _history_sort_key(record):
    dt = _parse_datetime_value(record.get("created_at"))
    return dt or datetime.min.replace(tzinfo=KST)


def _get_secret_value(key, default=""):
    try:
        return st.secrets.get(key, default)
    except Exception:
        return os.environ.get(key, default)


def _get_supabase_config():
    return {
        "url": _get_secret_value("SUPABASE_URL", ""),
        "key": _get_secret_value("SUPABASE_SERVICE_ROLE_KEY", "") or _get_secret_value("SUPABASE_KEY", ""),
        "bucket": _get_secret_value("SUPABASE_BUCKET", SUPABASE_HISTORY_BUCKET),
        "table": _get_secret_value("SUPABASE_TABLE", SUPABASE_HISTORY_TABLE),
        "settings_table": _get_secret_value("SUPABASE_SETTINGS_TABLE", SUPABASE_SETTINGS_TABLE),
    }


def _get_supabase_client():
    global _SUPABASE_CLIENT, _SUPABASE_ERROR
    config = _get_supabase_config()
    if not config["url"] or not config["key"]:
        missing = []
        if not config["url"]:
            missing.append("SUPABASE_URL")
        if not config["key"]:
            missing.append("SUPABASE_SERVICE_ROLE_KEY")
        _SUPABASE_ERROR = f"Secrets에 {', '.join(missing)} 값이 없습니다."
        return None
    if _SUPABASE_CLIENT is not None:
        return _SUPABASE_CLIENT
    try:
        from supabase import create_client
        _SUPABASE_CLIENT = create_client(config["url"], config["key"])
        _SUPABASE_ERROR = ""
        return _SUPABASE_CLIENT
    except ImportError:
        _SUPABASE_ERROR = "requirements.txt에 supabase가 설치되지 않았거나 배포가 다시 빌드되지 않았습니다."
        return None
    except Exception as e:
        _SUPABASE_ERROR = f"Supabase 연결 실패: {e}"
        return None


def _is_supabase_enabled():
    return _get_supabase_client() is not None


def _supabase_status_message():
    if _is_supabase_enabled():
        return "저장 위치: Supabase 공동 저장소"
    return f"저장 위치: 로컬 폴더 (임시 저장) · {_SUPABASE_ERROR}"


def _storage_path(history_id, file_name):
    folder = _safe_storage_segment(history_id, "history")
    name = _safe_storage_segment(file_name, "file")
    return f"{folder}/{name}"


def _upload_supabase_file(client, bucket, path, data, mime_type):
    try:
        return client.storage.from_(bucket).upload(
            path,
            data,
            file_options={"content-type": mime_type, "upsert": "true"},
        )
    except Exception:
        try:
            client.storage.from_(bucket).remove([path])
            return client.storage.from_(bucket).upload(
                path,
                data,
                file_options={"content-type": mime_type, "upsert": "true"},
            )
        except Exception as e:
            raise e


def _download_supabase_file(path):
    client = _get_supabase_client()
    if client is None:
        return None
    try:
        data = client.storage.from_(_get_supabase_config()["bucket"]).download(path)
        if isinstance(data, bytes):
            return data
        if hasattr(data, "content"):
            return data.content
        return bytes(data)
    except Exception:
        return None


def _list_review_history_local(limit=None):
    if not os.path.isdir(REVIEW_HISTORY_DIR):
        return []

    records = []
    for entry in os.listdir(REVIEW_HISTORY_DIR):
        record_dir = os.path.join(REVIEW_HISTORY_DIR, entry)
        meta_path = os.path.join(record_dir, "metadata.json")
        if not os.path.isdir(record_dir) or not os.path.exists(meta_path):
            continue
        meta = _read_json(meta_path, {})
        if not isinstance(meta, dict):
            continue
        meta["history_id"] = entry
        meta["history_dir"] = record_dir
        meta["storage_provider"] = "local"
        records.append(meta)

    records.sort(key=_history_sort_key, reverse=True)
    if limit:
        return records[:limit]
    return records


def _list_review_history_supabase(limit=None):
    client = _get_supabase_client()
    if client is None:
        return []
    table = _get_supabase_config()["table"]
    try:
        query = client.table(table).select("*").order("created_at", desc=True)
        if limit:
            query = query.limit(limit)
        response = query.execute()
    except Exception:
        return []

    records = []
    for row in response.data or []:
        metadata = row.get("metadata") or {}
        if not isinstance(metadata, dict):
            continue
        metadata["history_id"] = row.get("history_id") or metadata.get("history_id")
        metadata["created_at"] = row.get("created_at") or metadata.get("created_at")
        metadata["original_name"] = row.get("original_name") or metadata.get("original_name", "문서")
        metadata["correction_count"] = row.get("correction_count", metadata.get("correction_count", 0))
        metadata["storage_provider"] = "supabase"
        records.append(metadata)
    return records


def _list_review_history(limit=None):
    if _is_supabase_enabled():
        records_by_id = {}
        for record in _list_review_history_local():
            history_id = record.get("history_id")
            if history_id:
                records_by_id[history_id] = record
        for record in _list_review_history_supabase():
            history_id = record.get("history_id")
            if history_id:
                records_by_id[history_id] = record

        records = sorted(records_by_id.values(), key=_history_sort_key, reverse=True)
        if limit:
            return records[:limit]
        return records

    return _list_review_history_local(limit)


def _load_review_history_local(history_id):
    safe_id = _safe_filename(history_id)
    record_dir = os.path.join(REVIEW_HISTORY_DIR, safe_id)
    meta_path = os.path.join(record_dir, "metadata.json")
    meta = _read_json(meta_path, {})
    if not isinstance(meta, dict):
        return None
    meta["history_id"] = safe_id
    meta["history_dir"] = record_dir
    meta["storage_provider"] = "local"
    return meta


def _load_review_history_supabase(history_id):
    client = _get_supabase_client()
    if client is None:
        return None
    try:
        response = (
            client.table(_get_supabase_config()["table"])
            .select("*")
            .eq("history_id", _safe_filename(history_id))
            .limit(1)
            .execute()
        )
    except Exception:
        return None
    if not response.data:
        return None
    row = response.data[0]
    metadata = row.get("metadata") or {}
    if not isinstance(metadata, dict):
        return None
    metadata["history_id"] = row.get("history_id") or _safe_filename(history_id)
    metadata["created_at"] = row.get("created_at") or metadata.get("created_at")
    metadata["original_name"] = row.get("original_name") or metadata.get("original_name", "문서")
    metadata["correction_count"] = row.get("correction_count", metadata.get("correction_count", 0))
    metadata["storage_provider"] = "supabase"
    return metadata


def _load_review_history(history_id):
    if _is_supabase_enabled():
        return _load_review_history_supabase(history_id) or _load_review_history_local(history_id)
    return _load_review_history_local(history_id)


def _save_review_history_local(metadata, excel_data=None, completed_data=None):
    os.makedirs(REVIEW_HISTORY_DIR, exist_ok=True)

    history_id = st.session_state.get("current_review_history_id")
    if not history_id:
        timestamp = _now_kst().strftime("%Y%m%d_%H%M%S")
        name_part = os.path.splitext(_safe_filename(metadata.get("original_name", "문서")))[0][:40]
        history_id = f"{timestamp}_{name_part}"
        st.session_state.current_review_history_id = history_id

    record_dir = os.path.join(REVIEW_HISTORY_DIR, history_id)
    os.makedirs(record_dir, exist_ok=True)
    existing_metadata = _read_json(os.path.join(record_dir, "metadata.json"), {})
    if isinstance(existing_metadata, dict) and existing_metadata.get("created_at"):
        metadata["created_at"] = existing_metadata["created_at"]

    files = metadata.get("files", {})
    if excel_data:
        excel_name = _safe_filename(metadata.get("excel_name", "교정결과.xlsx"))
        with open(os.path.join(record_dir, excel_name), "wb") as f:
            f.write(excel_data)
        files["excel"] = excel_name

    if completed_data:
        completed_name = _safe_filename(metadata.get("completed_name", "완료본"))
        with open(os.path.join(record_dir, completed_name), "wb") as f:
            f.write(completed_data)
        files["completed"] = completed_name

    metadata["files"] = files
    metadata["history_id"] = history_id
    metadata["saved_at"] = _now_kst_iso()
    _write_json(os.path.join(record_dir, "metadata.json"), metadata)
    return history_id


def _save_review_history_supabase(metadata, excel_data=None, completed_data=None):
    client = _get_supabase_client()
    if client is None:
        return _save_review_history_local(metadata, excel_data, completed_data)

    config = _get_supabase_config()
    history_id = st.session_state.get("current_review_history_id")
    if not history_id:
        timestamp = _now_kst().strftime("%Y%m%d_%H%M%S")
        name_part = os.path.splitext(_safe_filename(metadata.get("original_name", "문서")))[0][:40]
        history_id = f"{timestamp}_{name_part}"
        st.session_state.current_review_history_id = history_id
    history_id = _safe_filename(history_id)

    existing_metadata = _load_review_history_supabase(history_id)
    if isinstance(existing_metadata, dict) and existing_metadata.get("created_at"):
        metadata["created_at"] = existing_metadata["created_at"]

    files = metadata.get("files", {})
    if excel_data:
        excel_name = _safe_filename(metadata.get("excel_name", "교정결과.xlsx"))
        excel_path = _storage_path(history_id, excel_name)
        _upload_supabase_file(
            client,
            config["bucket"],
            excel_path,
            excel_data,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        files["excel"] = excel_name
        files["excel_path"] = excel_path

    if completed_data:
        completed_name = _safe_filename(metadata.get("completed_name", "완료본"))
        completed_path = _storage_path(history_id, completed_name)
        _upload_supabase_file(
            client,
            config["bucket"],
            completed_path,
            completed_data,
            metadata.get("completed_mime", "application/octet-stream"),
        )
        files["completed"] = completed_name
        files["completed_path"] = completed_path

    metadata["files"] = files
    metadata["history_id"] = history_id
    metadata["storage_provider"] = "supabase"
    metadata["saved_at"] = _now_kst_iso()

    row = {
        "history_id": history_id,
        "created_at": metadata.get("created_at"),
        "saved_at": metadata.get("saved_at"),
        "original_name": metadata.get("original_name", "문서"),
        "correction_count": metadata.get("correction_count", 0),
        "metadata": metadata,
    }
    client.table(config["table"]).upsert(row, on_conflict="history_id").execute()
    return history_id


def _save_review_history(metadata, excel_data=None, completed_data=None):
    if _is_supabase_enabled():
        try:
            return _save_review_history_supabase(metadata, excel_data, completed_data)
        except Exception as e:
            st.warning(f"Supabase 저장에 실패해 로컬에 임시 저장했습니다: {e}")
    return _save_review_history_local(metadata, excel_data, completed_data)


def _load_persistent_json(setting_key, local_path, default=None):
    if default is None:
        default = {}

    client = _get_supabase_client()
    if client is not None:
        try:
            response = (
                client.table(_get_supabase_config()["settings_table"])
                .select("value")
                .eq("setting_key", setting_key)
                .limit(1)
                .execute()
            )
            if response.data:
                value = response.data[0].get("value")
                if isinstance(value, type(default)):
                    return value
                if isinstance(default, dict) and isinstance(value, dict):
                    return value
                if isinstance(default, list) and isinstance(value, list):
                    return value
        except Exception as e:
            st.warning(f"Supabase 설정 데이터 불러오기 실패({setting_key}): {e}")

    loaded = _read_json(local_path, default)
    if not isinstance(loaded, type(default)):
        return default

    if client is not None and loaded:
        try:
            client.table(_get_supabase_config()["settings_table"]).upsert(
                {
                    "setting_key": setting_key,
                    "value": loaded,
                    "saved_at": _now_kst_iso(),
                },
                on_conflict="setting_key",
            ).execute()
        except Exception:
            pass

    return loaded


def _save_persistent_json(setting_key, local_path, data):
    client = _get_supabase_client()
    if client is not None:
        try:
            client.table(_get_supabase_config()["settings_table"]).upsert(
                {
                    "setting_key": setting_key,
                    "value": data,
                    "saved_at": _now_kst_iso(),
                },
                on_conflict="setting_key",
            ).execute()
        except Exception as e:
            st.warning(f"Supabase 설정 데이터 저장 실패({setting_key}). 로컬에 임시 저장합니다: {e}")

    _write_json(local_path, data)


def _format_history_time(value):
    if not value:
        return ""
    dt = _parse_datetime_value(value)
    if dt:
        return dt.strftime("%Y-%m-%d %H:%M")
    return value


def _parse_history_datetime(value):
    return _parse_datetime_value(value)


def _set_loaded_history(history_id):
    st.session_state.loaded_history_id = history_id
    st.session_state.uploader_id = st.session_state.get("uploader_id", 0) + 1
    st.session_state.corrections = None
    st.session_state.script_text = None
    st.session_state.full_text = None
    st.session_state.score_result = None
    st.session_state.locations = None
    st.session_state.content_reviews = []
    st.session_state.alignment_report = None
    st.session_state.current_review_history_id = None


def _delete_review_history_local(history_id):
    safe_id = _safe_filename(history_id)
    record_dir = os.path.abspath(os.path.join(REVIEW_HISTORY_DIR, safe_id))
    history_root = os.path.abspath(REVIEW_HISTORY_DIR)
    if not record_dir.startswith(history_root + os.sep):
        return False
    if os.path.isdir(record_dir):
        shutil.rmtree(record_dir)
        if st.session_state.get("loaded_history_id") == safe_id:
            st.session_state.loaded_history_id = None
        return True
    return False


def _delete_review_history_supabase(history_id):
    client = _get_supabase_client()
    if client is None:
        return False
    safe_id = _safe_filename(history_id)
    record = _load_review_history_supabase(safe_id)
    files = (record or {}).get("files") or {}
    paths = [path for path in [files.get("excel_path"), files.get("completed_path")] if path]
    try:
        if paths:
            client.storage.from_(_get_supabase_config()["bucket"]).remove(paths)
        client.table(_get_supabase_config()["table"]).delete().eq("history_id", safe_id).execute()
        if st.session_state.get("loaded_history_id") == safe_id:
            st.session_state.loaded_history_id = None
        return True
    except Exception:
        return False


def _delete_review_history(history_id):
    if _is_supabase_enabled():
        deleted_supabase = _delete_review_history_supabase(history_id)
        deleted_local = _delete_review_history_local(history_id)
        return deleted_supabase or deleted_local
    return _delete_review_history_local(history_id)


def _filter_history_records(records, search_text, period_option, start_date=None, end_date=None):
    filtered = records
    query = (search_text or "").strip().lower()
    if query:
        filtered = [
            record for record in filtered
            if query in record.get("original_name", "").lower()
            or query in record.get("history_id", "").lower()
            or query in record.get("reference", "").lower()
            or query in record.get("knowledge", "").lower()
        ]

    now = _now_kst()
    if period_option == "오늘":
        start_dt = datetime(now.year, now.month, now.day, tzinfo=KST)
        filtered = [
            record for record in filtered
            if (dt := _parse_history_datetime(record.get("created_at"))) and dt >= start_dt
        ]
    elif period_option == "최근 7일":
        start_dt = now - timedelta(days=7)
        filtered = [
            record for record in filtered
            if (dt := _parse_history_datetime(record.get("created_at"))) and dt >= start_dt
        ]
    elif period_option == "최근 30일":
        start_dt = now - timedelta(days=30)
        filtered = [
            record for record in filtered
            if (dt := _parse_history_datetime(record.get("created_at"))) and dt >= start_dt
        ]
    elif period_option == "직접 지정" and start_date and end_date:
        start_dt = datetime.combine(start_date, datetime.min.time(), tzinfo=KST)
        end_dt = datetime.combine(end_date, datetime.max.time(), tzinfo=KST)
        filtered = [
            record for record in filtered
            if (dt := _parse_history_datetime(record.get("created_at"))) and start_dt <= dt <= end_dt
        ]

    return filtered


def _sort_history_records(records, sort_option):
    if sort_option == "오래된순":
        return sorted(records, key=_history_sort_key)
    if sort_option == "오류 많은 순":
        return sorted(records, key=lambda item: item.get("correction_count", 0), reverse=True)
    if sort_option == "파일명순":
        return sorted(records, key=lambda item: item.get("original_name", ""))
    return sorted(records, key=_history_sort_key, reverse=True)


def _render_history_manager_content():
    records = _list_review_history()
    st.markdown(
        """
        <style>
        .history-summary {
            color: #667085;
            font-size: 0.88rem;
            margin: -0.25rem 0 0.85rem 0;
        }
        .history-list-item {
            border: 1px solid #e6eaf0;
            border-radius: 8px;
            padding: 0.45rem 0.65rem;
            margin: 0.5rem 0 0.25rem 0;
            background: #ffffff;
        }
        .history-list-title {
            color: #101828;
            font-size: 0.96rem;
            font-weight: 700;
            line-height: 1.35;
            margin-bottom: 0.25rem;
            word-break: keep-all;
            overflow-wrap: anywhere;
        }
        .history-list-meta {
            color: #667085;
            font-size: 0.82rem;
            line-height: 1.45;
            overflow-wrap: anywhere;
        }
        .history-page-caption {
            color: #667085;
            font-size: 0.84rem;
            padding-top: 0.45rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(f'<div class="history-summary">전체 저장 기록 {len(records):,}개</div>', unsafe_allow_html=True)

    search_text = st.text_input("검색", placeholder="파일명, 지식명, 강의계획서명으로 검색", key="history_search_text")
    col_period, col_sort, col_page_size = st.columns([1.2, 1, 0.8])
    with col_period:
        period_option = st.selectbox(
            "기간 필터",
            ["전체", "오늘", "최근 7일", "최근 30일", "직접 지정"],
            key="history_period_option",
        )
    with col_sort:
        sort_option = st.selectbox(
            "정렬",
            ["최신순", "오래된순", "오류 많은 순", "파일명순"],
            key="history_sort_option",
        )
    with col_page_size:
        page_size = st.selectbox("페이지당", [5, 10, 20, 50], index=1, key="history_page_size")

    start_date = None
    end_date = None
    if period_option == "직접 지정":
        date_col1, date_col2 = st.columns(2)
        with date_col1:
            start_date = st.date_input("시작일", key="history_start_date")
        with date_col2:
            end_date = st.date_input("종료일", key="history_end_date")

    filtered_records = _filter_history_records(records, search_text, period_option, start_date, end_date)
    filtered_records = _sort_history_records(filtered_records, sort_option)

    total_pages = max(1, (len(filtered_records) + page_size - 1) // page_size)
    current_page = min(st.session_state.get("history_page", 1), total_pages)
    st.session_state.history_page = current_page

    nav_col1, nav_col2, nav_col3, nav_col4 = st.columns([1, 1, 2, 1])
    with nav_col1:
        if st.button("이전", disabled=current_page <= 1, key="history_prev_page"):
            st.session_state.history_page = max(1, current_page - 1)
            st.session_state.show_history_manager = True
            st.rerun()
    with nav_col2:
        if st.button("다음", disabled=current_page >= total_pages, key="history_next_page"):
            st.session_state.history_page = min(total_pages, current_page + 1)
            st.session_state.show_history_manager = True
            st.rerun()
    with nav_col3:
        st.markdown(
            f'<div class="history-page-caption">{current_page} / {total_pages} 페이지 · 검색 결과 {len(filtered_records):,}개</div>',
            unsafe_allow_html=True,
        )
    with nav_col4:
        st.checkbox("삭제 활성화", key="history_delete_enabled")

    delete_enabled = bool(st.session_state.get("history_delete_enabled", False))

    start_idx = (current_page - 1) * page_size
    page_records = filtered_records[start_idx:start_idx + page_size]

    if not page_records:
        st.info("조건에 맞는 검토 기록이 없습니다.")
        return

    for record in page_records:
        title = record.get("original_name", "문서")
        created = _format_history_time(record.get("created_at"))
        error_count = record.get("correction_count", 0)
        reference = record.get("reference", "사용 안함")
        knowledge = record.get("knowledge", "선택 안함")
        display_title = html.escape(str(title))
        display_created = html.escape(str(created))
        display_reference = html.escape(str(reference))
        display_knowledge = html.escape(str(knowledge))

        with st.container():
            st.markdown('<div class="history-list-item">', unsafe_allow_html=True)
            row_info, row_view, row_delete = st.columns([5.2, 1.35, 0.9], vertical_alignment="center")
            with row_info:
                st.markdown(
                    f"""
                    <div class="history-list-title">{display_created} · {display_title}</div>
                    <div class="history-list-meta">오류 {error_count}건 · 지식: {display_knowledge} · 강의계획서: {display_reference}</div>
                    """,
                    unsafe_allow_html=True,
                )
            with row_view:
                if st.button("열기", key=f"manager_load_{record['history_id']}", use_container_width=True):
                    st.session_state.show_history_manager = False
                    _set_loaded_history(record["history_id"])
                    st.rerun()
            with row_delete:
                if st.button("삭제", key=f"manager_delete_{record['history_id']}", disabled=not delete_enabled, use_container_width=True):
                    if _delete_review_history(record["history_id"]):
                        st.success("삭제되었습니다.")
                        st.session_state.show_history_manager = True
                        st.rerun()
                    else:
                        st.error("삭제하지 못했습니다.")
            st.markdown('</div>', unsafe_allow_html=True)


def _show_history_manager():
    if hasattr(st, "dialog"):
        try:
            dialog_decorator = st.dialog("검토 기록 전체 보기", width="large")
        except TypeError:
            dialog_decorator = st.dialog("검토 기록 전체 보기")

        @dialog_decorator
        def _history_dialog():
            if st.button("닫기", key="close_history_dialog"):
                st.session_state.show_history_manager = False
                st.rerun()
            _render_history_manager_content()

        _history_dialog()
    else:
        with st.expander("검토 기록 전체 보기", expanded=True):
            if st.button("닫기", key="close_history_panel"):
                st.session_state.show_history_manager = False
                st.rerun()
            _render_history_manager_content()


def _build_correction_rows(corrections, locations, file_ext):
    rows = []
    label_map = {'spelling': '맞춤법/오타', 'foreign': '외래어 표기', 'spacing': '띄어쓰기'}
    for old, new in (corrections or {}).items():
        locs = (locations or {}).get(old, [])
        loc_str = ", ".join(map(str, locs))
        if loc_str:
            if file_ext == '.pdf':
                loc_str += " 페이지"
            elif file_ext == '.pptx':
                loc_str += " 슬라이드"
        rows.append({
            "발생 위치": loc_str,
            "수정 전(원본)": old,
            "수정 후(AI 제안)": new,
            "오류 유형": label_map.get(core.classify_error(old, new), '기타')
        })
    return rows


def _render_saved_history_record(record):
    st.subheader("📚 저장된 검토 결과")
    st.caption(f"{record.get('original_name', '문서')} · {_format_history_time(record.get('created_at'))}")

    score_result = record.get("score_result")
    if score_result:
        st.markdown("#### 🏅 문서 품질 점수")
        render_score_dashboard(score_result)

    alignment_report = record.get("alignment_report")
    if alignment_report:
        st.markdown("#### 📘 강의계획서 기준 내용 적합성 분석")
        st.metric("내용 적합성", f"{alignment_report.get('overall_score', 0)}점")
        if alignment_report.get("verdict"):
            st.write(alignment_report.get("verdict", ""))
        if alignment_report.get("summary"):
            st.info(alignment_report["summary"])
        gaps = alignment_report.get("gaps") or []
        if gaps:
            gap_rows = []
            for gap in gaps:
                if isinstance(gap, dict):
                    gap_rows.append({
                        "검토 기준": gap.get("criterion", ""),
                        "강의계획서 근거": gap.get("reference_evidence", ""),
                        "현재 문서 상태": gap.get("document_evidence", ""),
                        "보완 제안": gap.get("recommendation", ""),
                    })
            if gap_rows:
                st.dataframe(pd.DataFrame(gap_rows), use_container_width=True, hide_index=True)

    content_reviews = record.get("content_reviews") or []
    if content_reviews:
        st.markdown("#### 🔎 위치별 내용 오류 검토")
        st.dataframe(pd.DataFrame([
            {
                "발생 위치": review.get("location", ""),
                "오류 유형": review.get("issue_type", ""),
                "문제가 된 원문": review.get("original_excerpt", ""),
                "검토 의견": review.get("issue", ""),
                "판단 근거": review.get("basis", ""),
                "수정 제안": review.get("recommendation", ""),
            }
            for review in content_reviews if isinstance(review, dict)
        ]), use_container_width=True, hide_index=True)

    correction_rows = record.get("correction_rows") or _build_correction_rows(
        record.get("corrections", {}),
        record.get("locations", {}),
        record.get("file_ext", ""),
    )
    st.markdown("#### 📋 수정 전 / 수정 후 검토")
    if correction_rows:
        st.dataframe(pd.DataFrame(correction_rows), use_container_width=True, hide_index=True)
    else:
        st.info("AI가 변경할 곳을 찾지 못했던 문서입니다.")

    files = record.get("files") or {}
    record_dir = record.get("history_dir", "")
    if files:
        st.markdown("#### 📥 저장된 파일 다운로드")
        col_excel, col_completed = st.columns(2)
        excel_file = files.get("excel")
        completed_file = files.get("completed")
        if excel_file:
            excel_data = None
            if record.get("storage_provider") == "supabase" and files.get("excel_path"):
                excel_data = _download_supabase_file(files["excel_path"])
            else:
                excel_path = os.path.join(record_dir, excel_file)
                if os.path.exists(excel_path):
                    with open(excel_path, "rb") as f:
                        excel_data = f.read()
            if excel_data:
                col_excel.download_button(
                    "📊 교정 결과 엑셀 다운로드",
                    data=excel_data,
                    file_name=excel_file,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
        if completed_file:
            completed_data = None
            if record.get("storage_provider") == "supabase" and files.get("completed_path"):
                completed_data = _download_supabase_file(files["completed_path"])
            else:
                completed_path = os.path.join(record_dir, completed_file)
                if os.path.exists(completed_path):
                    with open(completed_path, "rb") as f:
                        completed_data = f.read()
            if completed_data:
                col_completed.download_button(
                    "💖 저장된 완성본 다운로드",
                    data=completed_data,
                    file_name=completed_file,
                    mime=record.get("completed_mime", "application/octet-stream"),
                    use_container_width=True,
                )


# 로고 (사이드바 열림/닫힘 모두 표시)
_logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ARASoft로고.png")
if os.path.exists(_logo_path):
    st.markdown(
        """
        <style>
        .st-emotion-cache-4xtz07 {
            height: 3rem !important;
            margin-top: 2.25rem !important;
        }
        .st-emotion-cache-1h1td79 hr {
            margin: 0rem 0px !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    try:
        st.logo(_logo_path, size="large")
    except TypeError:
        # size 파라미터 미지원 버전
        st.logo(_logo_path)
    except AttributeError:
        # st.logo 자체 미지원 구버전 fallback
        _logo_b64 = base64.b64encode(open(_logo_path, "rb").read()).decode()
        st.markdown(
            f"""
            <style>
            [data-testid="stToolbar"]::before {{
                content: "";
                display: inline-block;
                background-image: url("data:image/png;base64,{_logo_b64}");
                background-size: contain;
                background-repeat: no-repeat;
                background-position: center;
                width: 120px;
                height: 34px;
                vertical-align: middle;
                margin-right: 8px;
            }}
            </style>
            """,
            unsafe_allow_html=True,
        )

# 사이드바

with st.sidebar:
    selected_model = "gpt-5.6-sol"
    open_history_manager = False

    st.divider()
    st.subheader("📚 검토 완료 문서")
    history_records = _list_review_history()
    st.caption(f"총 {len(history_records):,}개 저장됨")
    if history_records:
        for record in history_records[:5]:
            title = record.get("original_name", "문서")
            created = _format_history_time(record.get("created_at"))
            error_count = record.get("correction_count", 0)
            button_label = f"{created} · {title[:18]} · {error_count}건"
            if st.button(button_label, key=f"load_history_{record['history_id']}", use_container_width=True):
                _set_loaded_history(record["history_id"])
                st.rerun()
        if st.button("📂 리스트 전체 보기", use_container_width=True):
            open_history_manager = True
            st.session_state.show_history_manager = True
    else:
        st.caption("아직 저장된 검토 결과가 없습니다.")

    if st.session_state.get("loaded_history_id"):
        if st.button("불러온 검토 닫기", use_container_width=True):
            st.session_state.loaded_history_id = None
            st.rerun()

    st.divider()
    st.subheader("🧠 AI 사전 학습 (지식 베이스)")
    
    kb_file_path = "knowledge_base.json"
    knowledge_base = _load_persistent_json("knowledge_base", kb_file_path, {})

    def _save_knowledge_base(data):
        _save_persistent_json("knowledge_base", kb_file_path, data)
                
    new_keyword = st.text_input("학습할 주제/키워드 명 (예: 소형무인기논문)", placeholder="키워드 입력")
    kb_file = st.file_uploader("학습할 문서 업로드 (선택, PPTX/PDF/TXT/DOCX)", type=["pptx", "pdf", "txt", "docx"])
    
    if st.button("🚀 지식 학습 시작"):
        target_keyword = new_keyword.strip()
        if kb_file and not target_keyword:
            target_keyword = os.path.splitext(kb_file.name)[0]
            
        if not target_keyword:
            st.error("주제명(키워드)을 입력하거나 문서를 업로드해주세요.")
        elif not API_KEY_DEFAULT or not API_KEY_DEFAULT.startswith("sk-"):
            st.error("OpenAI API 키가 설정되어 있지 않습니다.")
        else:
            with st.spinner(f"'{target_keyword}'에 대한 전문 지식 생성 중..."):
                kb_data = None
                if kb_file:
                    ext = os.path.splitext(kb_file.name)[1].lower()
                    file_text = ""
                    try:
                        if ext == ".pdf":
                            import fitz
                            doc = fitz.open(stream=kb_file.read(), filetype="pdf")
                            file_text = core.extract_full_text_pdf(doc)
                        elif ext == ".pptx":
                            doc = Presentation(kb_file)
                            file_text = core.extract_full_text_pptx(doc)
                        elif ext == ".txt":
                            file_text = kb_file.read().decode("utf-8", errors="ignore")
                        elif ext == ".docx":
                            import docx
                            doc = docx.Document(kb_file)
                            file_text = "\n".join([p.text for p in doc.paragraphs])
                    except ImportError:
                        if ext == ".docx":
                            st.error("Word(.docx) 처리를 위해 python-docx 패키지가 필요합니다.")
                        else:
                            st.error("모듈 불러오기 실패.")
                    except Exception as e:
                        st.error(f"파일 읽기 오류: {e}")
                        
                    if file_text.strip():
                        kb_data = core.generate_knowledge_from_text(file_text, API_KEY_DEFAULT, model="gpt-5.6-sol")
                    else:
                        st.error("문서에서 텍스트를 추출하지 못했습니다.")
                else:
                    kb_data = core.generate_knowledge(target_keyword, API_KEY_DEFAULT, model="gpt-5.6-sol")
                    
                if kb_data:
                    knowledge_base[target_keyword] = kb_data
                    _save_knowledge_base(knowledge_base)
                    st.success(f"'{target_keyword}' 학습 완료 및 저장됨!")
                else:
                    if kb_file and file_text.strip():
                        st.error("지식 생성에 실패했습니다.")
                    elif not kb_file:
                        st.error("지식 생성에 실패했습니다.")
    
    if knowledge_base:
        with st.expander("📚 현재 학습된 지식 목록 보기", expanded=False):
            for kw in list(knowledge_base.keys()):
                if st.session_state.get(f"edit_mode_{kw}", False):
                    new_name = st.text_input("새 이름", value=kw, key=f"new_name_{kw}", label_visibility="collapsed")
                    col_s1, col_s2, col_s3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col_s2:
                        if st.button("💾", key=f"save_{kw}", help="저장", type="tertiary"):
                            if new_name and new_name != kw:
                                knowledge_base[new_name] = knowledge_base.pop(kw)
                                _save_knowledge_base(knowledge_base)
                            st.session_state[f"edit_mode_{kw}"] = False
                            st.rerun()
                    with col_s3:
                        if st.button("❌", key=f"cancel_{kw}", help="취소", type="tertiary"):
                            st.session_state[f"edit_mode_{kw}"] = False
                            st.rerun()
                else:
                    col1, col2, col3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col1:
                        st.caption(f"- {kw} ({len(knowledge_base[kw].get('terms', []))}개 용어)")
                    with col2:
                        if st.button("✏️", key=f"edit_{kw}", help=f"'{kw}' 이름 수정", type="tertiary"):
                            st.session_state[f"edit_mode_{kw}"] = True
                            st.rerun()
                    with col3:
                        if st.button("🗑️", key=f"del_{kw}", help=f"'{kw}' 지식 삭제", type="tertiary"):
                            del knowledge_base[kw]
                            _save_knowledge_base(knowledge_base)
                            st.rerun()

    st.divider()
    st.subheader("📘 강의계획서 관리")

    reference_documents_path = "reference_documents.json"
    reference_documents = _load_persistent_json("reference_documents", reference_documents_path, {})

    def _save_reference_documents(documents):
        _save_persistent_json("reference_documents", reference_documents_path, documents)

    new_reference_name = st.text_input(
        "강의계획서 이름",
        placeholder="예: 2026 데이터분석 입문"
    )
    new_reference_file = st.file_uploader(
        "새 강의계획서 업로드",
        type=["hwp", "hwpx", "docx", "pdf", "pptx", "txt"],
        key="new_reference_document_uploader"
    )

    if st.button("➕ 강의계획서 등록"):
        reference_name = new_reference_name.strip()
        if new_reference_file and not reference_name:
            reference_name = os.path.splitext(new_reference_file.name)[0]

        if not new_reference_file:
            st.error("등록할 강의계획서 파일을 올려주세요.")
        elif not reference_name:
            st.error("강의계획서 이름을 입력해주세요.")
        else:
            try:
                new_reference_file.seek(0)
                extracted_text = core.extract_reference_document_text(
                    new_reference_file.name,
                    new_reference_file.read()
                )
                if not extracted_text.strip():
                    st.error("문서에서 텍스트를 추출하지 못했습니다. DOCX 또는 PDF로 변환해 다시 시도해주세요.")
                else:
                    reference_documents[reference_name] = {
                        "source_filename": new_reference_file.name,
                        "text": extracted_text,
                    }
                    _save_reference_documents(reference_documents)
                    st.success(f"'{reference_name}' 강의계획서가 저장되었습니다.")
                    st.rerun()
            except Exception as e:
                st.error(f"강의계획서 등록 오류: {e}")

    if reference_documents:
        with st.expander("📘 저장된 강의계획서 목록", expanded=False):
            for reference_name in list(reference_documents.keys()):
                reference_data = reference_documents[reference_name]
                reference_length = len(reference_data.get("text", "")) if isinstance(reference_data, dict) else len(str(reference_data))

                if st.session_state.get(f"edit_reference_mode_{reference_name}", False):
                    renamed_reference = st.text_input(
                        "새 이름",
                        value=reference_name,
                        key=f"new_reference_name_{reference_name}",
                        label_visibility="collapsed"
                    )
                    col_s1, col_s2, col_s3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col_s2:
                        if st.button("💾", key=f"save_reference_{reference_name}", help="저장", type="tertiary"):
                            renamed_reference = renamed_reference.strip()
                            if renamed_reference and renamed_reference != reference_name:
                                reference_documents[renamed_reference] = reference_documents.pop(reference_name)
                                _save_reference_documents(reference_documents)
                            st.session_state[f"edit_reference_mode_{reference_name}"] = False
                            st.rerun()
                    with col_s3:
                        if st.button("❌", key=f"cancel_reference_{reference_name}", help="취소", type="tertiary"):
                            st.session_state[f"edit_reference_mode_{reference_name}"] = False
                            st.rerun()
                else:
                    col1, col2, col3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col1:
                        st.caption(f"- {reference_name} ({reference_length:,}자)")
                    with col2:
                        if st.button("✏️", key=f"edit_reference_{reference_name}", help=f"'{reference_name}' 이름 수정", type="tertiary"):
                            st.session_state[f"edit_reference_mode_{reference_name}"] = True
                            st.rerun()
                    with col3:
                        if st.button("🗑️", key=f"delete_reference_{reference_name}", help=f"'{reference_name}' 삭제", type="tertiary"):
                            reference_documents.pop(reference_name)
                            _save_reference_documents(reference_documents)
                            st.rerun()

    st.divider()
    st.subheader("📖 사용자 맞춤법 사전")
    
    sp_dict_file_path = "custom_spelling_dicts.json"
    spelling_dicts = _load_persistent_json("custom_spelling_dicts", sp_dict_file_path, {})
    
    def _save_all_spelling_dicts(dicts_to_save):
        _save_persistent_json("custom_spelling_dicts", sp_dict_file_path, dicts_to_save)
        # 하위 호환성을 위해 모든 사전의 단어를 맞춤법사전.txt에 통합 저장
        all_words = []
        for words in dicts_to_save.values():
            all_words.extend(words)
        unique_words = sorted(list(set(all_words)))
        with open("맞춤법사전.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(unique_words))

    # 데이터 로드 및 마이그레이션
    if not spelling_dicts:
        # 기존 맞춤법사전.txt가 있으면 가져와서 '기본 사전'으로 마이그레이션
        old_dict_path = "맞춤법사전.txt"
        if os.path.exists(old_dict_path):
            try:
                with open(old_dict_path, "r", encoding="utf-8") as f:
                    old_text = f.read()
                raw_words = old_text.replace('\n', ',').split(',')
                words_list = [w.strip() for w in raw_words if w.strip()]
                if words_list:
                    spelling_dicts["기본 사전"] = words_list
                    _save_all_spelling_dicts(spelling_dicts)
            except Exception:
                pass

    new_dict_name = st.text_input("새 맞춤법 사전 이름", placeholder="예: IT 용어 사전")
    new_dict_words = st.text_area(
        "예외 단어 입력 (쉼표(,)나 줄바꿈으로 구분)",
        height=100,
        placeholder="단어1\n단어2"
    )

    if st.button("➕ 맞춤법 사전 등록"):
        target_name = new_dict_name.strip()
        if not target_name:
            st.error("사전 이름을 입력해주세요.")
        else:
            raw_w = new_dict_words.replace('\n', ',').split(',')
            w_list = [w.strip() for w in raw_w if w.strip()]
            spelling_dicts[target_name] = w_list
            _save_all_spelling_dicts(spelling_dicts)
            st.success(f"'{target_name}' 사전 등록 완료!")
            st.rerun()

    if spelling_dicts:
        with st.expander("📖 등록된 맞춤법 사전 목록 보기", expanded=False):
            for dn in list(spelling_dicts.keys()):
                words_str = "\n".join(spelling_dicts[dn])

                if st.session_state.get(f"edit_sp_mode_{dn}", False):
                    new_dn = st.text_input("새 사전 이름", value=dn, key=f"new_dn_{dn}", label_visibility="collapsed")
                    new_words_val = st.text_area("단어 편집", value=words_str, key=f"edit_words_{dn}", height=120)

                    col_s1, col_s2, col_s3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col_s2:
                        if st.button("💾", key=f"save_sp_{dn}", help="저장", type="tertiary"):
                            raw_w = new_words_val.replace('\n', ',').split(',')
                            w_list = [w.strip() for w in raw_w if w.strip()]
                            if new_dn and new_dn != dn:
                                spelling_dicts.pop(dn)
                                spelling_dicts[new_dn] = w_list
                            else:
                                spelling_dicts[dn] = w_list
                            _save_all_spelling_dicts(spelling_dicts)
                            st.session_state[f"edit_sp_mode_{dn}"] = False
                            st.rerun()
                    with col_s3:
                        if st.button("❌", key=f"cancel_sp_{dn}", help="취소", type="tertiary"):
                            st.session_state[f"edit_sp_mode_{dn}"] = False
                            st.rerun()
                else:
                    col1, col2, col3 = st.columns([7.5, 1.2, 1.3], vertical_alignment="center")
                    with col1:
                        st.caption(f"- {dn} ({len(spelling_dicts[dn])}개 단어)")
                    with col2:
                        if st.button("✏️", key=f"edit_sp_{dn}", help=f"'{dn}' 이름 및 단어 수정", type="tertiary"):
                            st.session_state[f"edit_sp_mode_{dn}"] = True
                            st.rerun()
                    with col3:
                        if st.button("🗑️", key=f"del_sp_{dn}", help=f"'{dn}' 사전 삭제", type="tertiary"):
                            spelling_dicts.pop(dn)
                            _save_all_spelling_dicts(spelling_dicts)
                            st.rerun()

# ==========================================
# 점수 대시보드 렌더링 함수
# ==========================================
def render_score_dashboard(sr):
    """가중치 기반 문서 품질 점수 대시보드를 HTML로 렌더링한다."""
    score = sr['score']
    grade_label = sr['grade_label']
    grade_color = sr['grade_color']
    total_words = sr['total_words']
    total_errors = sr['total_errors']
    ec = sr['error_counts']
    wsum = sr['weighted_error_sum']

    # 점수에 따른 배경 그라디언트 색상
    bg_start = "#1a1a2e"
    bg_end   = "#16213e"

    html = f"""
    <style>
      @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;700;900&display=swap');
      .score-card {{
        font-family: 'Inter', sans-serif;
        background: linear-gradient(135deg, {bg_start} 0%, {bg_end} 100%);
        border-radius: 18px;
        padding: 28px 32px;
        margin: 8px 0 20px 0;
        border: 1px solid rgba(255,255,255,0.1);
        box-shadow: 0 8px 40px rgba(0,0,0,0.4);
      }}
      .score-title {{
        color: rgba(255,255,255,0.65);
        font-size: 15px;
        font-weight: 700;
        letter-spacing: 1px;
        text-transform: uppercase;
        margin: 0 0 18px 0;
      }}
      .score-main {{
        display: flex;
        align-items: center;
        gap: 28px;
        margin-bottom: 22px;
      }}
      .score-number {{
        font-size: 80px;
        font-weight: 900;
        color: {grade_color};
        line-height: 1;
        text-shadow: 0 0 30px {grade_color}66;
        min-width: 140px;
        text-align: center;
      }}
      .score-unit {{
        color: rgba(255,255,255,0.4);
        font-size: 16px;
        text-align: center;
        margin-top: 4px;
      }}
      .score-right {{ flex: 1; }}
      .grade-label {{
        font-size: 24px;
        font-weight: 700;
        color: {grade_color};
        margin-bottom: 14px;
      }}
      .bar-bg {{
        background: rgba(255,255,255,0.12);
        border-radius: 10px;
        height: 14px;
        overflow: hidden;
      }}
      .bar-fill {{
        background: linear-gradient(90deg, {grade_color}99, {grade_color});
        height: 100%;
        width: {score}%;
        border-radius: 10px;
      }}
      .bar-labels {{
        display: flex;
        justify-content: space-between;
        color: rgba(255,255,255,0.3);
        font-size: 11px;
        margin-top: 5px;
      }}
      .stat-grid {{
        display: grid;
        grid-template-columns: repeat(4, 1fr);
        gap: 12px;
        margin-bottom: 16px;
      }}
      .stat-box {{
        background: rgba(255,255,255,0.05);
        border-radius: 12px;
        padding: 14px 10px;
        text-align: center;
        border: 1px solid rgba(255,255,255,0.07);
      }}
      .stat-label {{
        color: rgba(255,255,255,0.45);
        font-size: 11px;
        margin-bottom: 6px;
        letter-spacing: 0.3px;
      }}
      .stat-value {{
        font-size: 22px;
        font-weight: 800;
      }}
      .stat-weight {{
        color: rgba(255,255,255,0.25);
        font-size: 10px;
        margin-top: 3px;
      }}
      .score-footnote {{
        padding: 10px 14px;
        background: rgba(255,255,255,0.03);
        border-radius: 8px;
        border-left: 3px solid {grade_color};
        color: rgba(255,255,255,0.4);
        font-size: 11.5px;
        line-height: 1.6;
      }}
    </style>

    <div class="score-card">
      <p class="score-title">📊 문서 품질 점수 (가중치 기반)</p>

      <div class="score-main">
        <div>
          <div class="score-number">{score}</div>
          <div class="score-unit">/ 100점</div>
        </div>
        <div class="score-right">
          <div class="grade-label">{grade_label}</div>
          <div class="bar-bg"><div class="bar-fill"></div></div>
          <div class="bar-labels"><span>0</span><span>50</span><span>100</span></div>
        </div>
      </div>

      <div class="stat-grid">
        <div class="stat-box">
          <div class="stat-label">총 어절 수</div>
          <div class="stat-value" style="color:#fff;">{total_words:,}</div>
          <div class="stat-weight">분석 대상</div>
        </div>
        <div class="stat-box">
          <div class="stat-label">맞춤법 / 오타</div>
          <div class="stat-value" style="color:#E74C3C;">{ec['spelling']}건</div>
          <div class="stat-weight">가중치 ×2.0 (심각)</div>
        </div>
        <div class="stat-box">
          <div class="stat-label">외래어 표기</div>
          <div class="stat-value" style="color:#F39C12;">{ec['foreign']}건</div>
          <div class="stat-weight">가중치 ×1.5 (보통)</div>
        </div>
        <div class="stat-box">
          <div class="stat-label">띄어쓰기</div>
          <div class="stat-value" style="color:#3498DB;">{ec['spacing']}건</div>
          <div class="stat-weight">가중치 ×1.0 (경미)</div>
        </div>
      </div>

      <div class="score-footnote">
        💡 <b>점수 산출 공식</b>: (1 − 가중 오류 합계 / 총 어절 수) × 100 &nbsp;|&nbsp;
        가중 오류 합계: <b>{wsum}</b>점 &nbsp;|&nbsp;
        실제 오류 발생 횟수 합계: <b>{total_errors}건</b><br>
        맞춤법·오타는 2배, 외래어 표기는 1.5배, 띄어쓰기는 1배로 감점됩니다.
      </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)


# ==========================================
# 등급별 색상 안내 범례
# ==========================================
def render_grade_legend():
    grades = [
        ("S", "95~100점", "#FFD700", "🏆 최우수"),
        ("A", "85~94점",  "#2ECC71", "✅ 우수"),
        ("B", "70~84점",  "#3498DB", "🔵 양호"),
        ("C", "50~69점",  "#F39C12", "⚠️ 미흡"),
        ("D", "0~49점",   "#E74C3C", "🔴 불량"),
    ]
    cols = st.columns(5)
    for col, (g, rng, color, label) in zip(cols, grades):
        col.markdown(
            f"<div style='text-align:center; background:rgba(255,255,255,0.05);"
            f"border-radius:10px; padding:10px 4px; border:1px solid {color}44;'>"
            f"<div style='font-size:22px; font-weight:900; color:{color};'>{g}</div>"
            f"<div style='font-size:12px; color:rgba(255,255,255,0.6);'>{rng}</div>"
            f"<div style='font-size:11px; color:{color}; margin-top:3px;'>{label}</div>"
            f"</div>",
            unsafe_allow_html=True
        )


# 메인 영역
if open_history_manager or st.session_state.get("show_history_manager"):
    _show_history_manager()

st.markdown(
    """
    <div style='background-color: rgba(128, 128, 128, 0.08); padding: 15px; border-radius: 10px; border-left: 5px solid #FF00E5; margin-bottom: 20px; color: inherit;'>
        <h4 style='margin-top: 0; color: #FF00E5;'>💡 파일 업로드 가이드</h4>
        <ul style='margin-bottom: 0; padding-left: 20px; font-size: 15px; line-height: 1.6;'>
            <li>한글 파일 : DOCX 파일로 변화 하여 업로드하는 것을 추천(한글 파일 업로드 시 DOCX 파일로 변환되어 추출되나 표 등 깨짐 현상 있음)</li>
            <li>PDF 파일 : PDF로 추출되나 수정은 안되고 수정해야 될 부분이 체크되서 추출됨</li>
            <li>PPT 파일 : 교정된 형태로 PPT 추출됨</li>
        </ul>
    </div>
    """,
    unsafe_allow_html=True
)

st.subheader("📁 1. 파일 업로드")

# 지식 및 사전 선택 (가로 배치)
col_kb, col_sp = st.columns(2)
with col_kb:
    kb_options = ["선택 안함"] + list(knowledge_base.keys()) if 'knowledge_base' in locals() else ["선택 안함"]
    selected_kb_keyword = st.selectbox("검사에 적용할 사전 학습 지식 (선택)", options=kb_options)

with col_sp:
    # 맞춤법 사전 선택 (다중 선택)
    sp_options = list(spelling_dicts.keys()) if 'spelling_dicts' in locals() else []
    selected_sp_dicts = st.multiselect(
        "검사에 적용할 사용자 맞춤법 사전 (다중 선택)",
        options=sp_options
    )

st.markdown("#### 📘 강의계획서 기준 분석 (선택)")
reference_options = ["사용 안함"] + list(reference_documents.keys()) if 'reference_documents' in locals() else ["사용 안함"]
col_reference, col_lesson = st.columns([2, 1])
with col_reference:
    selected_reference_name = st.selectbox(
        "분석에 적용할 강의계획서",
        options=reference_options,
        help="새 강의계획서는 왼쪽 사이드바의 '강의계획서 관리'에서 한 번만 등록하면 됩니다."
    )
with col_lesson:
    lesson_options = ["전체/자동"] + [f"{number}차시" for number in range(1, 13)]
    selected_lesson_label = st.selectbox(
        "검토할 차시",
        options=lesson_options,
        help="특정 차시를 선택하면 강의계획서에서 해당 차시의 학습목표와 내용 범위를 우선 검토합니다."
    )
reference_text = ""
if selected_reference_name != "사용 안함":
    selected_reference_data = reference_documents.get(selected_reference_name, {})
    if isinstance(selected_reference_data, dict):
        reference_text = selected_reference_data.get("text", "")
    else:
        reference_text = str(selected_reference_data)

    if reference_text.strip():
        st.success(f"'{selected_reference_name}' 강의계획서 적용 준비 완료")
        with st.expander("저장된 강의계획서 내용 미리보기", expanded=False):
            st.text(reference_text[:3000] + ("\n..." if len(reference_text) > 3000 else ""))
    else:
        st.warning("저장된 강의계획서 내용이 비어 있습니다. 사이드바에서 다시 등록해주세요.")

# 엑셀 이미지 포함 옵션 추가
export_images = st.checkbox("엑셀 다운로드용 슬라이드/페이지 이미지 추출 (LibreOffice/PowerPoint COM 작동, 수십 초 소요)", value=False, help="체크하면 엑셀 파일에 슬라이드 이미지가 삽입되지만, 검사 속도가 느려집니다. 체크 해제 시 이미지 없이 빠르게 다운로드 가능합니다.")


if "uploader_id" not in st.session_state:
    st.session_state.uploader_id = 0

uploaded_file = st.file_uploader(
    "검사할 문서를 올려주세요.", 
    type=["pptx", "hwp", "hwpx", "docx", "pdf"],
    key=f"file_uploader_{st.session_state.uploader_id}"
)

if uploaded_file is not None:
    st.success(f"'{uploaded_file.name}' 업로드 성공!")
    
    # 세션 상태 초기화
    for key in [
        'corrections', 'script_text', 'full_text', 'score_result', 'alignment_report',
        'alignment_reference_name', 'alignment_lesson_label', 'content_reviews',
        'content_review_context'
    ]:
        if key not in st.session_state:
            st.session_state[key] = None
        
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    # 업로드된 파일을 메모리 기반 객체로 로드
    doc_obj = None
    hwp_text_content = ""
    
    # 파일 포인터를 처음으로 돌려줍니다. (BadZipFile 에러 예방)
    uploaded_file.seek(0)
    
    if file_ext == '.pdf':
        import fitz
        file_bytes = uploaded_file.read()
        doc_obj = fitz.open(stream=file_bytes, filetype="pdf")
    elif file_ext == '.pptx':
        doc_obj = Presentation(uploaded_file)
    elif file_ext == '.docx':
        import docx
        doc_obj = docx.Document(uploaded_file)
    elif file_ext == '.hwp':
        file_bytes = uploaded_file.read()
        hwp_text_content = core.extract_text_hwp(file_bytes)
        doc_obj = file_bytes
    elif file_ext == '.hwpx':
        file_bytes = uploaded_file.read()
        hwp_text_content = core.extract_text_hwpx(file_bytes)
        doc_obj = file_bytes
        
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("🚀 AI 분석 및 텍스트 스캔 시작", use_container_width=True):
            # 이전 결과 초기화
            st.session_state.loaded_history_id = None
            st.session_state.current_review_history_id = None
            st.session_state.corrections = None
            st.session_state.script_text = None
            st.session_state.full_text = None
            st.session_state.score_result = None
            st.session_state.alignment_report = None
            st.session_state.alignment_reference_name = None
            st.session_state.alignment_lesson_label = None
            st.session_state.content_reviews = []
            st.session_state.content_review_context = None
            st.session_state.img_cache = {}  # 이미지 캐시 초기화
            
            with st.spinner("문서를 스캔하고 대본을 추출하는 중..."):
                if file_ext == '.pdf':
                    script_text = core.extract_narrations_pdf(doc_obj)
                    full_text   = core.extract_full_text_pdf(doc_obj)
                elif file_ext == '.pptx':
                    script_text = core.extract_narrations(doc_obj)
                    full_text   = core.extract_full_text_pptx(doc_obj)
                elif file_ext == '.docx':
                    script_text = {}
                    full_text   = core.extract_full_text_docx(doc_obj)
                elif file_ext in ('.hwp', '.hwpx'):
                    script_text = {}
                    full_text   = hwp_text_content
                st.session_state.script_text = script_text
                st.session_state.full_text   = full_text
                
            st.success(f"대본 추출 완료! 이제 문서 검사에 진입합니다.")
            
            if not API_KEY_DEFAULT or not API_KEY_DEFAULT.startswith("sk-"):
                st.error("서버에 올바른 OpenAI API 환경변수 비밀키가 설정되어 있지 않습니다!")
                st.stop()
                
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            def update_progress(current, total):
                progress = int((current / total) * 100)
                progress_bar.progress(progress)
                status_text.markdown(f"**진행 상황 (맞춤법 스캔):** {current}/{total} 페이지/슬라이드 스캔 완료... ({selected_model} 사용 중)")
            
            # 선택된 맞춤법 사전들로부터 단어 취합
            custom_dict_list = []
            if 'selected_sp_dicts' in locals() and selected_sp_dicts:
                for dn in selected_sp_dicts:
                    custom_dict_list.extend(spelling_dicts.get(dn, []))
            
            # 선택된 지식 베이스가 있다면 용어 목록을 맞춤법 예외 사전에 병합
            active_kb_data = None
            if selected_kb_keyword != "선택 안함" and 'knowledge_base' in locals():
                active_kb_data = knowledge_base.get(selected_kb_keyword)
                if active_kb_data:
                    kb_terms = active_kb_data.get("terms", [])
                    custom_dict_list.extend(kb_terms)
                    
            with st.spinner(f"OpenAI 맞춤법 스캔 중 (1단계) ({selected_model})..."):
                if file_ext == '.pdf':
                    corrections, locations = core.get_openai_corrections_by_page_pdf(
                        doc_obj, 
                        API_KEY_DEFAULT, 
                        is_paid_tier=True,
                        custom_dict=custom_dict_list,
                        progress_callback=update_progress,
                        model=selected_model
                    )
                elif file_ext == '.pptx':
                    corrections, locations = core.get_openai_corrections_by_slide(
                        doc_obj, 
                        API_KEY_DEFAULT, 
                        is_paid_tier=True,
                        custom_dict=custom_dict_list,
                        progress_callback=update_progress,
                        model=selected_model
                    )
                elif file_ext == '.docx':
                    corrections, locations = core.get_openai_corrections_docx(
                        doc_obj,
                        API_KEY_DEFAULT,
                        is_paid_tier=True,
                        custom_dict=custom_dict_list,
                        progress_callback=update_progress,
                        model=selected_model
                    )
                elif file_ext in ('.hwp', '.hwpx'):
                    corrections, locations = core.get_openai_corrections_hwp_text(
                        full_text,
                        API_KEY_DEFAULT,
                        is_paid_tier=True,
                        custom_dict=custom_dict_list,
                        progress_callback=update_progress,
                        model=selected_model
                    )
                st.session_state.corrections = corrections
                st.session_state.locations = locations

            # 강의계획서 기준 분석은 기존 사전 학습 및 맞춤법 검사와 별도로 수행한다.
            if reference_text.strip():
                with st.spinner(f"강의계획서 기준 내용 적합성 분석 중 (2단계) ({selected_model})..."):
                    try:
                        st.session_state.alignment_report = core.analyze_document_against_reference(
                            st.session_state.full_text,
                            reference_text,
                            API_KEY_DEFAULT,
                            document_name=uploaded_file.name,
                            reference_name=selected_reference_name,
                            lesson_label=selected_lesson_label,
                            model=selected_model,
                        )
                        st.session_state.alignment_reference_name = selected_reference_name
                        st.session_state.alignment_lesson_label = selected_lesson_label
                    except Exception as e:
                        st.session_state.alignment_report = None
                        st.session_state.alignment_reference_name = None
                        st.session_state.alignment_lesson_label = None
                        st.error(f"강의계획서 기준 분석 중 오류가 발생했습니다: {e}")

            # 모든 문서 형식에 대해 AI 사전 지식과 강의계획서를 함께 사용해 위치별 내용 오류를 찾는다.
            if active_kb_data or reference_text.strip():
                with st.spinner(f"AI 사전 지식·강의계획서 기반 위치별 내용 검토 중 ({selected_model})..."):
                    document_sections = core.build_document_content_sections(
                        file_ext,
                        doc_obj=doc_obj,
                        full_text=st.session_state.full_text or "",
                    )
                    progress_bar_rev = st.progress(0)
                    status_text_rev = st.empty()

                    def update_progress_rev(current, total):
                        progress = int((current / total) * 100) if total else 100
                        progress_bar_rev.progress(progress)
                        status_text_rev.markdown(
                            f"**진행 상황 (내용 검토):** {current}/{total}개 구간 검토 완료... ({selected_model} 사용 중)"
                        )

                    try:
                        st.session_state.content_reviews = core.get_openai_content_reviews_by_sections(
                            document_sections,
                            active_kb_data,
                            API_KEY_DEFAULT,
                            reference_text=reference_text,
                            lesson_label=selected_lesson_label,
                            progress_callback=update_progress_rev,
                            model=selected_model,
                        )
                        st.session_state.content_review_context = {
                            "knowledge": selected_kb_keyword,
                            "reference": selected_reference_name,
                            "lesson": selected_lesson_label,
                        }
                        progress_bar_rev.progress(100)
                        status_text_rev.markdown("**✅ 위치별 내용 검토 완료!**")
                    except Exception as e:
                        st.session_state.content_reviews = []
                        st.session_state.content_review_context = None
                        status_text_rev.empty()
                        st.error(f"위치별 내용 검토 중 오류가 발생했습니다: {e}")

            # ── 점수 계산 ──────────────────────────────────
            if st.session_state.full_text:
                st.session_state.score_result = core.calculate_score(
                    corrections,
                    st.session_state.full_text
                )
                
            progress_bar.progress(100)
            status_text.markdown("**✅ AI 분석 완료!**")

    # ──────────────────────────────────────────────
    # 점수 대시보드 표시
    # ──────────────────────────────────────────────
    if st.session_state.score_result is not None:
        if st.button("🔄 검사 결과 초기화 (새 파일 올리기)", use_container_width=True):
            st.session_state.uploader_id += 1
            st.session_state.corrections = None
            st.session_state.script_text = None
            st.session_state.full_text = None
            st.session_state.score_result = None
            st.session_state.locations = None
            st.session_state.content_reviews = []
            st.session_state.alignment_report = None
            st.session_state.alignment_reference_name = None
            st.session_state.alignment_lesson_label = None
            st.session_state.content_review_context = None
            st.session_state.img_cache = {}  # 이미지 캐시 초기화
            st.rerun()

        st.subheader("🏅 문서 품질 점수")
        render_score_dashboard(st.session_state.score_result)
        with st.expander("📘 등급 기준표 보기"):
            render_grade_legend()

    alignment_report = st.session_state.get('alignment_report')
    if (
        alignment_report
        and st.session_state.get('alignment_reference_name') == selected_reference_name
        and st.session_state.get('alignment_lesson_label') == selected_lesson_label
    ):
        st.subheader("📘 강의계획서 기준 내용 적합성 분석")
        st.caption(f"검토 범위: {selected_lesson_label}")
        score_col, verdict_col = st.columns([1, 4])
        with score_col:
            st.metric("내용 적합성", f"{alignment_report.get('overall_score', 0)}점")
        with verdict_col:
            st.markdown("**종합 판정**")
            st.write(alignment_report.get('verdict', ''))

        if alignment_report.get('summary'):
            st.info(alignment_report['summary'])

        strengths = alignment_report.get('strengths') or []
        if strengths:
            with st.expander("✅ 잘 반영된 내용", expanded=True):
                for item in strengths:
                    st.write(f"- {item}")

        gaps = alignment_report.get('gaps') or []
        if gaps:
            st.markdown("**누락·불일치 및 보완 사항**")
            gap_rows = []
            for gap in gaps:
                if isinstance(gap, dict):
                    gap_rows.append({
                        "검토 기준": gap.get("criterion", ""),
                        "강의계획서 근거": gap.get("reference_evidence", ""),
                        "현재 문서 상태": gap.get("document_evidence", ""),
                        "보완 제안": gap.get("recommendation", ""),
                    })
                else:
                    gap_rows.append({"검토 기준": str(gap), "강의계획서 근거": "", "현재 문서 상태": "", "보완 제안": ""})
            st.dataframe(pd.DataFrame(gap_rows), use_container_width=True, hide_index=True)
        else:
            st.success("강의계획서 기준에서 뚜렷한 누락이나 불일치가 발견되지 않았습니다.")

        recommendations = alignment_report.get('recommendations') or []
        if recommendations:
            with st.expander("🛠 우선 수정 권고", expanded=True):
                for idx, item in enumerate(recommendations, 1):
                    st.write(f"{idx}. {item}")

    current_content_context = {
        "knowledge": selected_kb_keyword,
        "reference": selected_reference_name,
        "lesson": selected_lesson_label,
    }
    content_reviews = st.session_state.get('content_reviews') or []
    if st.session_state.get('content_review_context') == current_content_context:
        st.subheader("🔎 위치별 내용 오류 검토")
        st.caption(f"검토 범위: {selected_lesson_label}")
        if content_reviews:
            content_review_df = pd.DataFrame([
                {
                    "발생 위치": review.get("location", ""),
                    "오류 유형": review.get("issue_type", ""),
                    "문제가 된 원문": review.get("original_excerpt", ""),
                    "검토 의견": review.get("issue", ""),
                    "판단 근거": review.get("basis", ""),
                    "수정 제안": review.get("recommendation", ""),
                }
                for review in content_reviews if isinstance(review, dict)
            ])
            st.dataframe(content_review_df, use_container_width=True, hide_index=True)
        else:
            st.success("선택한 지식과 차시 기준에서 뚜렷한 내용 오류가 발견되지 않았습니다.")

    if st.session_state.corrections is not None:
        st.subheader("📋 2. 수정 전 / 수정 후 검토")
        
        c_dict = st.session_state.corrections
        loc_dict = st.session_state.get('locations', {})
        excel_data = None
        correction_rows_for_history = _build_correction_rows(c_dict, loc_dict, file_ext)
        if len(c_dict) == 0:
            st.info("AI가 변경할 곳을 찾지 못했습니다. 문장이 이미 완벽하거나 수정할 내용이 없습니다.")
        else:
            # 엑셀용 이미지 추출 (미리 캐싱)
            img_cache = {}
            if export_images:
                if "img_cache" not in st.session_state or st.session_state.img_cache is None:
                    st.session_state.img_cache = {}
                
                unique_locs = set()
                for old in c_dict.keys():
                    locs = loc_dict.get(old, [])
                    if locs:
                        unique_locs.add(locs[0])
                
                # 캐시되지 않은 위치의 이미지만 추출하여 캐싱
                missing_locs = [loc for loc in unique_locs if loc not in st.session_state.img_cache]
                
                if missing_locs:
                    with st.spinner("엑셀 다운로드를 위한 원본 이미지 준비 중 (수 초 소요될 수 있습니다)..."):
                        if file_ext == '.pdf':
                            for loc in missing_locs:
                                st.session_state.img_cache[loc] = core.get_pdf_page_image_bytes(doc_obj, loc)
                        elif file_ext == '.pptx':
                            uploaded_file.seek(0)
                            pptx_bytes = uploaded_file.read()
                            new_imgs = core.get_pptx_slide_images(pptx_bytes, missing_locs)
                            st.session_state.img_cache.update(new_imgs)
                
                img_cache = st.session_state.img_cache

            # 오류 유형 컬럼 추가
            rows = []
            image_mappings = []
            seen_locs = set()
            
            for row in correction_rows_for_history:
                old = row["수정 전(원본)"]
                locs = loc_dict.get(old, [])
                loc_str = row["발생 위치"]
                
                if loc_str and loc_str not in seen_locs:
                    img_bytes = img_cache.get(locs[0]) if locs else None
                    seen_locs.add(loc_str)
                else:
                    img_bytes = None
                
                image_mappings.append(img_bytes)
                
                rows.append({
                    "발생 위치": row["발생 위치"],
                    "원본 이미지": "",
                    "수정 전(원본)": row["수정 전(원본)"],
                    "수정 후(AI 제안)": row["수정 후(AI 제안)"],
                    "오류 유형": row["오류 유형"]
                })
            df = pd.DataFrame(rows)
            # 화면에는 이미지 컬럼을 빼고 보여줌
            st.dataframe(df.drop(columns=["원본 이미지"]), use_container_width=True, hide_index=True)
            
            # 엑셀 다운로드 버튼 추가
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df.to_excel(writer, index=False, sheet_name='교정결과')
                workbook = writer.book
                worksheet = writer.sheets['교정결과']
                
                # 서식 정의
                bg_color_1 = '#FFFFFF'
                bg_color_2 = '#F4F8FC'
                
                fmt_c1 = workbook.add_format({'bg_color': bg_color_1, 'text_wrap': True, 'valign': 'vcenter', 'align': 'center', 'border': 1, 'border_color': '#D9D9D9'})
                fmt_l1 = workbook.add_format({'bg_color': bg_color_1, 'text_wrap': True, 'valign': 'vcenter', 'align': 'left', 'border': 1, 'border_color': '#D9D9D9'})
                
                fmt_c2 = workbook.add_format({'bg_color': bg_color_2, 'text_wrap': True, 'valign': 'vcenter', 'align': 'center', 'border': 1, 'border_color': '#D9D9D9'})
                fmt_l2 = workbook.add_format({'bg_color': bg_color_2, 'text_wrap': True, 'valign': 'vcenter', 'align': 'left', 'border': 1, 'border_color': '#D9D9D9'})
                
                # 열 너비 설정
                worksheet.set_column('A:A', 14)
                worksheet.set_column('B:B', 60)
                worksheet.set_column('C:C', 40)
                worksheet.set_column('D:D', 40)
                worksheet.set_column('E:E', 15)
                
                # 그룹별로 서식 적용 및 병합
                groups = []
                current_loc = None
                start_idx = 0
                for i in range(len(df)):
                    loc = df.iloc[i, 0] # 발생 위치
                    if loc != current_loc:
                        if current_loc is not None:
                            groups.append((start_idx, i - 1, current_loc))
                        current_loc = loc
                        start_idx = i
                if len(df) > 0:
                    groups.append((start_idx, len(df) - 1, current_loc))
                    
                img_scale = 0.17 if file_ext == '.pdf' else 0.32
                
                for group_idx, (s_idx, e_idx, loc_str) in enumerate(groups):
                    group_size = e_idx - s_idx + 1
                    is_even = (group_idx % 2 == 0)
                    fc = fmt_c1 if is_even else fmt_c2
                    fl = fmt_l1 if is_even else fmt_l2
                    
                    # 행 높이: 이미지가 잘리지 않도록 단일 항목일 때는 높이를 충분히 크게(190), 여러 개일 때는 골고루 분배
                    if group_size == 1:
                        row_h = 190
                    else:
                        row_h = max(190 // group_size, 45)
                        
                    for r in range(s_idx, e_idx + 1):
                        worksheet.set_row(r + 1, row_h)
                        
                    # 발생 위치, 원본 이미지 병합 (A, B열)
                    if group_size > 1:
                        worksheet.merge_range(s_idx + 1, 0, e_idx + 1, 0, loc_str, fc)
                        worksheet.merge_range(s_idx + 1, 1, e_idx + 1, 1, "", fc)
                    else:
                        worksheet.write(s_idx + 1, 0, loc_str, fc)
                        worksheet.write(s_idx + 1, 1, "", fc)
                        
                    # 이미지 삽입
                    img_bytes = image_mappings[s_idx]
                    if img_bytes:
                        # 이미지는 병합된 블록의 시작 셀(s_idx + 1)에 삽입
                        worksheet.insert_image(s_idx + 1, 1, f"img_{s_idx}.png", {
                            'image_data': io.BytesIO(img_bytes),
                            'x_scale': img_scale,
                            'y_scale': img_scale,
                            'x_offset': 5,
                            'y_offset': 5,
                            'object_position': 1
                        })
                        
                    # 나머지 열 데이터 쓰기 (C, D, E)
                    for r in range(s_idx, e_idx + 1):
                        worksheet.write(r + 1, 2, df.iloc[r, 2], fl) # 수정 전
                        worksheet.write(r + 1, 3, df.iloc[r, 3], fl) # 수정 후
                        worksheet.write(r + 1, 4, df.iloc[r, 4], fc) # 오류 유형
                
                # 내용 검토 시트 추가
                if st.session_state.get('content_reviews'):
                    review_df = pd.DataFrame([
                        {
                            "발생 위치": review.get("location", ""),
                            "오류 유형": review.get("issue_type", ""),
                            "문제가 된 원문": review.get("original_excerpt", ""),
                            "검토 의견": review.get("issue", ""),
                            "판단 근거": review.get("basis", ""),
                            "수정 제안": review.get("recommendation", ""),
                        }
                        for review in st.session_state.content_reviews
                        if isinstance(review, dict)
                    ])
                    review_df.to_excel(writer, index=False, sheet_name='내용검토')
                    review_ws = writer.sheets['내용검토']
                    review_ws.set_column('A:A', 20)
                    review_ws.set_column('B:B', 16)
                    review_ws.set_column('C:F', 45)
                    for r in range(len(review_df)):
                        review_ws.set_row(r + 1, 60)
                        review_ws.write(r + 1, 0, review_df.iloc[r, 0], fmt_c1)
                        review_ws.write(r + 1, 1, review_df.iloc[r, 1], fmt_c1)
                        for col_idx in range(2, 6):
                            review_ws.write(r + 1, col_idx, review_df.iloc[r, col_idx], fmt_l1)
                        
            excel_data = output.getvalue()
            
            st.download_button(
                label="📊 교정 결과 엑셀 다운로드",
                data=excel_data,
                file_name=f"교정결과_{uploaded_file.name}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
            
            if file_ext == '.pdf':
                st.warning("위 변경 사항들은 완성본 다운로드 시 '핑크색(FF00E5) 형광펜 (메모 코멘트)' 형태로 PDF에 표시됩니다.")
            elif file_ext == '.pptx':
                st.warning("위 변경 사항들은 완성본 다운로드 시 '핑크색(FF00E5)' 서식으로 PPT에 일괄 덮어씌워집니다. "
                           "(부분 굵게/색상 등 일부 인라인 서식은 초기화될 수 있습니다.)")
            elif file_ext == '.docx':
                st.warning("위 변경 사항들은 완성본 다운로드 시 '핑크색(FF00E5)' 서식으로 워드(Word) 파일에 덮어씌워집니다.")
            elif file_ext in ('.hwp', '.hwpx'):
                st.warning("위 변경 사항들은 완성본 다운로드 시 교정된 내용이 반영된 워드(.docx) 문서 파일로 자동 변환되어 다운로드됩니다.")
            
        st.subheader("📥 3. 완성본 다운로드")
        
        download_data = None
        completion_error = None
        mime_type = "application/octet-stream"
        btn_label = "💖 교정 반영본 다운로드"
        download_name = f"완료_{uploaded_file.name}"

        try:
            with st.spinner("수정 및 덧그리기 작업 중입니다..."):
                out_stream = io.BytesIO()
                if file_ext == '.pdf':
                    core.apply_corrections_to_pdf(doc_obj, st.session_state.corrections)
                    doc_obj.save(out_stream)
                    doc_obj.close()
                    mime_type = "application/pdf"
                    btn_label = "💖 교정 하이라이트 PDF 다운로드"
                    download_name = f"완료_{uploaded_file.name}"
                elif file_ext == '.pptx':
                    core.apply_corrections_to_ppt(doc_obj, st.session_state.corrections)
                    doc_obj.save(out_stream)
                    mime_type = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
                    btn_label = "💖 핑크색 교정 반영본 PPTX 다운로드"
                    download_name = f"완료_{uploaded_file.name}"
                elif file_ext == '.docx':
                    core.apply_corrections_to_docx(doc_obj, st.session_state.corrections)
                    doc_obj.save(out_stream)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    btn_label = "💖 핑크색 교정 반영본 DOCX 다운로드"
                    download_name = f"완료_{uploaded_file.name}"
                elif file_ext in ('.hwp', '.hwpx'):
                    try:
                        full_txt = st.session_state.full_text or ""
                        corrections = st.session_state.corrections or {}
                        # 한글 본문 텍스트와 교정 사전을 바탕으로 Word 문서(.docx) 생성
                        docx_doc = core.create_docx_from_hwp_text(full_txt, corrections)
                        docx_doc.save(out_stream)
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        btn_label = "💖 워드(.docx)로 변환된 교정 반영본 다운로드"
                        base_name = os.path.splitext(uploaded_file.name)[0]
                        download_name = f"완료_{base_name}.docx"
                    except Exception:
                        # 예외 발생 시 최종 백업으로 텍스트 파일 제공
                        corrected_text = core.apply_corrections_to_text(st.session_state.full_text or "", st.session_state.corrections or {})
                        out_stream.write(corrected_text.encode('utf-8'))
                        mime_type = "text/plain"
                        btn_label = "💖 교정 반영본 텍스트 파일(TXT) 다운로드 (대체)"
                        base_name = os.path.splitext(uploaded_file.name)[0]
                        download_name = f"완료_{base_name}.txt"

                download_data = out_stream.getvalue()
        except Exception as e:
            completion_error = str(e)
            st.error(f"완성본 파일 생성 중 오류가 발생했습니다. 검토 기록은 저장을 시도합니다: {completion_error}")

        if download_data:
            st.download_button(
                label=btn_label,
                data=download_data,
                file_name=download_name,
                mime=mime_type,
                use_container_width=True
            )

        review_metadata = {
            "created_at": _now_kst_iso(),
            "original_name": uploaded_file.name,
            "file_ext": file_ext,
            "selected_model": selected_model,
            "knowledge": selected_kb_keyword,
            "spelling_dicts": selected_sp_dicts,
            "reference": selected_reference_name,
            "lesson": selected_lesson_label,
            "score_result": st.session_state.get("score_result"),
            "correction_count": len(c_dict),
            "corrections": c_dict,
            "locations": loc_dict,
            "correction_rows": correction_rows_for_history,
            "alignment_report": st.session_state.get("alignment_report"),
            "content_reviews": st.session_state.get("content_reviews") or [],
            "content_review_context": st.session_state.get("content_review_context"),
            "excel_name": f"교정결과_{uploaded_file.name}.xlsx",
            "completed_name": download_name,
            "completed_mime": mime_type,
            "completed_error": completion_error,
        }
        saved_history_id = _save_review_history(
            review_metadata,
            excel_data=excel_data,
            completed_data=download_data,
        )
        st.caption(f"검토 결과가 저장되었습니다. 사이드바 '검토 완료 문서'에서 다시 열 수 있습니다. (ID: {saved_history_id})")

elif st.session_state.get("loaded_history_id"):
    loaded_record = _load_review_history(st.session_state.loaded_history_id)
    if loaded_record:
        _render_saved_history_record(loaded_record)
    else:
        st.warning("저장된 검토 결과를 불러오지 못했습니다.")

